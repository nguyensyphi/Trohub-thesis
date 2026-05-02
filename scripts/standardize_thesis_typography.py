from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


INPUT = Path(r"C:\Users\LENOVO\Downloads\ĐỒ ÁN\trohub\ReportThesis_NSPhi_final_ver_1_formatted_final.docx")
OUTPUT = Path(r"C:\Users\LENOVO\Downloads\ĐỒ ÁN\trohub\ReportThesis_NSPhi_final_ver_1_formatted_uniform.docx")


def set_east_asia_font(target, font_name: str) -> None:
    rpr = target._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def clear_run_spacing(run) -> None:
    rpr = run._element.get_or_add_rPr()
    spacing = rpr.find(qn("w:spacing"))
    if spacing is not None:
        rpr.remove(spacing)


def apply_run_font(run, font_name: str, size_pt: float, bold=None, italic=None) -> None:
    run.font.name = font_name
    set_east_asia_font(run, font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    clear_run_spacing(run)


def style_by_name(doc: Document, name: str):
    return doc.styles[name]


def configure_styles(doc: Document) -> None:
    normal = style_by_name(doc, "Normal")
    normal.font.name = "Times New Roman"
    set_east_asia_font(normal, "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size, bold, align in [
        ("Heading 1", 16, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 13, True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 4", 13, True, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = style_by_name(doc, name)
        style.font.name = "Times New Roman"
        set_east_asia_font(style, "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(6 if name != "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.alignment = align

    for style_name in ("TOC 1", "TOC 2", "TOC 3"):
        try:
            style = style_by_name(doc, style_name)
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        set_east_asia_font(style, "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def first_nonempty_index(paragraphs):
    for idx, p in enumerate(paragraphs):
        if p.text.strip():
            return idx
    return 0


def find_index(paragraphs, target: str):
    for idx, p in enumerate(paragraphs):
        if p.text.strip() == target:
            return idx
    return -1


def main() -> None:
    doc = Document(str(INPUT))
    configure_styles(doc)

    paragraphs = list(doc.paragraphs)
    first_heading_1 = next((i for i, p in enumerate(paragraphs) if p.style.name == "Heading 1"), len(paragraphs))
    references_idx = find_index(paragraphs, "REFERENCES")
    list_figures_idx = find_index(paragraphs, "LIST OF FIGURES")
    list_tables_idx = find_index(paragraphs, "LIST OF TABLES")
    abstract_idx = find_index(paragraphs, "ABSTRACT")

    for idx, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        is_caption = text.startswith("Figure ") or text.startswith("Table ")
        is_toc = p.style.name.startswith("TOC")
        is_cover = idx < first_heading_1
        in_lists = (
            (list_figures_idx != -1 and list_tables_idx != -1 and list_figures_idx < idx < list_tables_idx)
            or (list_tables_idx != -1 and abstract_idx != -1 and list_tables_idx < idx < abstract_idx)
        )
        in_references = references_idx != -1 and idx > references_idx

        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        if p.style.name == "Normal":
            if is_cover:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                pf.line_spacing = 1.0
            elif in_lists or is_toc or is_caption:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if in_lists or is_toc else WD_ALIGN_PARAGRAPH.CENTER
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                pf.line_spacing = 1.0
            elif in_references:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                pf.line_spacing = 1.5
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                pf.line_spacing = 1.5

        elif p.style.name == "Heading 1":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        size = 13
        bold = None
        italic = None

        if p.style.name == "Heading 1":
            size = 16
            bold = True
        elif p.style.name == "Heading 2":
            size = 14
            bold = True
        elif p.style.name in {"Heading 3", "Heading 4"}:
            size = 13
            bold = True
        elif is_caption:
            size = 12
        elif in_lists or is_toc:
            size = 13
        elif is_cover:
            size = 13

        if idx in {2, 9}:  # thesis title lines on the two cover pages
            size = 16
            bold = True

        for run in p.runs:
            if not run.text:
                continue
            apply_run_font(run, "Times New Roman", size, bold=bold, italic=italic)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.paragraph_format.line_spacing = 1.0
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        if run.text:
                            apply_run_font(run, "Times New Roman", 12)

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    main()
