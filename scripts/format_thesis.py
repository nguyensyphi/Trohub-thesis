import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


OUTPUT = Path(r"C:\Users\LENOVO\Downloads\ĐỒ ÁN\trohub\ReportThesis_NSPhi_final_ver_1_formatted_final.docx")
SOURCE = Path(r"C:\Users\LENOVO\Downloads\ĐỒ ÁN\trohub\ReportThesis_NSPhi_final_ver_1.docx")


FRONT_HEADINGS = {
    "ACKNOWLEDGMENTS",
    "LIST OF FIGURES",
    "LIST OF TABLES",
    "ABSTRACT",
    "REFERENCES",
}


def get_text(paragraph):
    return paragraph.text.strip()


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def normalize_text(text):
    t = text.replace("\t", " ")
    t = re.sub(r"^\s*#{1,6}\s*", "", t)
    t = t.replace("**", "")
    t = re.sub(r"^\s*\*\s+", "• ", t)
    t = re.sub(r"^\s*\s*", "• ", t)
    t = t.replace("―", '"').replace("‖", '"')
    t = t.replace("MOMO", "MoMo")
    t = t.replace("Password recover", "Password recovery")
    t = t.replace("Password recoveryy", "Password recovery")
    t = t.replace("Sent Email", "Send Email")
    t = re.sub(r"\s+,", ",", t)
    t = re.sub(r"\s+\.", ".", t)
    t = re.sub(r"\s+;", ";", t)
    t = re.sub(r"\s+:", ":", t)
    t = re.sub(r" {2,}", " ", t).strip()
    t = re.sub(r"(?<!\.)\.\.(?!\.)", ".", t)
    t = re.sub(r"the degree of\s+Bachelor", "the degree of Bachelor", t)

    exact = {
        "CHAPTER 4. IMPLEMENTS AND RESULTS": "CHAPTER 4. IMPLEMENTATION AND RESULTS",
        "4.1. Implement": "4.1. Implementation",
        "4.2. Result": "4.2. Results",
        "2.3. Summarize": "2.3. Summary",
        "2.3.1. Login and Forgot Password": "4.2.1.7. Login and Forgot Password",
        "4.2.3.2.": "",
        "Figure 16: Database diagram number 8 Users": "Figure 16: Database diagram number 9 Users",
        "APPROVED BY: ________________________________ ,": "APPROVED BY: ________________________________",
    }
    t = exact.get(t, t)
    t = re.sub(r"^4\.1\.\s*Implement$", "4.1. Implementation", t)
    return t


def heading_style(text):
    if text in FRONT_HEADINGS:
        return "Heading 1"
    if re.match(r"^CHAPTER\s+\d+\.\s+\S", text):
        return "Heading 1"
    if re.match(r"^UC-\d+:\s+\S", text):
        return "Heading 4"
    match = re.match(r"^(?P<num>\d+(?:\.\d+)+\.)\s+\S", text)
    if not match:
        return None
    depth = len(match.group("num").rstrip(".").split("."))
    return {
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
        5: "Heading 5",
        6: "Heading 6",
        7: "Heading 7",
        8: "Heading 8",
    }.get(depth, "Heading 9")


def parse_list_pages(section_text, label):
    pattern = re.compile(
        rf"({label}\s+(\d+):\s+.*?)(?:\s*\.{{3,}}\s*|\t|\s{{2,}})(\d+)(?=\s+{label}\s+\d+:|$)",
        re.IGNORECASE | re.DOTALL,
    )
    pages = {}
    for _, number, page in pattern.findall(section_text):
        pages[int(number)] = page
    return pages


def clean_section_blob(paragraphs):
    return " ".join(" ".join(p.text.split()) for p in paragraphs if p.text.strip())


def collect_body_captions(paragraphs, start_index):
    figures = {}
    tables = {}
    for paragraph in paragraphs[start_index:]:
        text = normalize_text(get_text(paragraph))
        fig = re.match(r"^Figure\s+(\d+):\s+(.+)$", text)
        if fig:
            figures[int(fig.group(1))] = f"Figure {int(fig.group(1))}: {fig.group(2)}"
            continue
        tab = re.match(r"^Table\s+(\d+):\s+(.+)$", text)
        if tab:
            tables[int(tab.group(1))] = f"Table {int(tab.group(1))}: {tab.group(2)}"
    return figures, tables


def rebuild_list_section(heading_paragraph, next_heading_paragraph, entries):
    current = heading_paragraph._p.getnext()
    stop = next_heading_paragraph._p
    while current is not None and current is not stop:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    anchor = heading_paragraph
    for entry in entries:
        para = insert_paragraph_after(anchor, entry, "Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.tab_stops.clear_all()
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        para.paragraph_format.space_before = 0
        para.paragraph_format.space_after = 0
        anchor = para


doc = Document(str(SOURCE))
paragraphs = list(doc.paragraphs)

lof_idx = next(i for i, p in enumerate(paragraphs) if get_text(p) == "LIST OF FIGURES")
lot_idx = next(i for i, p in enumerate(paragraphs) if get_text(p) == "LIST OF TABLES")
abstract_idx = next(i for i, p in enumerate(paragraphs) if get_text(p) == "ABSTRACT")

figure_pages = parse_list_pages(clean_section_blob(paragraphs[lof_idx + 1 : lot_idx]), "Figure")
table_pages = parse_list_pages(clean_section_blob(paragraphs[lot_idx + 1 : abstract_idx]), "Table")

for idx, paragraph in enumerate(paragraphs):
    style_name = paragraph.style.name
    if style_name.startswith("TOC"):
        continue

    if lof_idx < idx < lot_idx or lot_idx < idx < abstract_idx:
        continue

    raw_text = paragraph.text
    text = get_text(paragraph)
    if not text:
        if raw_text and not raw_text.strip():
            delete_paragraph(paragraph)
        continue

    new_text = normalize_text(text)
    if new_text in {"", "."}:
        delete_paragraph(paragraph)
        continue

    if new_text != text:
        paragraph.text = new_text

    new_text = get_text(paragraph)

    if re.match(r"^(Figure|Table)\s+\d+:\s+\S", new_text):
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    target_style = heading_style(new_text)
    if target_style:
        paragraph.style = doc.styles[target_style]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif paragraph.style.name.startswith("Heading"):
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


paragraphs = list(doc.paragraphs)
lof_para = next(p for p in paragraphs if get_text(p) == "LIST OF FIGURES")
lot_para = next(p for p in paragraphs if get_text(p) == "LIST OF TABLES")
abstract_para = next(p for p in paragraphs if get_text(p) == "ABSTRACT")

paragraphs = list(doc.paragraphs)
abstract_idx = next(i for i, p in enumerate(paragraphs) if get_text(p) == "ABSTRACT")
body_figures, body_tables = collect_body_captions(paragraphs, abstract_idx + 1)

figure_entries = []
for number in sorted(body_figures):
    page = figure_pages.get(number, "")
    figure_entries.append(f"{body_figures[number]}\t{page}".rstrip())

table_entries = []
for number in sorted(body_tables):
    page = table_pages.get(number, "")
    table_entries.append(f"{body_tables[number]}\t{page}".rstrip())

rebuild_list_section(lof_para, lot_para, figure_entries)
rebuild_list_section(lot_para, abstract_para, table_entries)

doc.save(str(OUTPUT))
